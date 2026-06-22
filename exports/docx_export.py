"""
DOCX meeting report generation using python-docx.

Same content as the PDF report, but as an editable Word document — useful
for teams who want to tweak wording before sharing internally.
"""

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, RGBColor

from utils.logger import get_logger

log = get_logger(__name__)


def _add_heading(doc: Document, text: str):
    heading = doc.add_heading(text, level=2)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)


def export_summary_docx(analysis: dict, meeting_insights: dict, output_path: str, context: str = "") -> str:
    """Writes a formatted DOCX report and returns the output path."""
    doc = Document()

    title = doc.add_heading("Meeting Summary Report", level=0)

    if context:
        p = doc.add_paragraph()
        run = p.add_run(f"Context: {context}")
        run.italic = True

    # --- TL;DR ---
    _add_heading(doc, "TL;DR")
    doc.add_paragraph(analysis.get("tldr") or "Not available.")

    # --- Key Decisions ---
    _add_heading(doc, "Key Decisions")
    decisions = analysis.get("key_decisions") or []
    if decisions:
        for d in decisions:
            doc.add_paragraph(d, style="List Bullet")
    else:
        doc.add_paragraph("None identified.")

    # --- Action Items table ---
    _add_heading(doc, "Action Items")
    action_items = analysis.get("action_items") or []
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Light Grid Accent 1"
    header_cells = table.rows[0].cells
    for i, label in enumerate(["Task", "Owner", "Deadline", "Priority"]):
        header_cells[i].text = label
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    if action_items:
        for item in action_items:
            row_cells = table.add_row().cells
            row_cells[0].text = item.get("task") or "Not specified"
            row_cells[1].text = item.get("owner") or "Not specified"
            row_cells[2].text = item.get("deadline") or "Not specified"
            row_cells[3].text = item.get("priority") or "Medium"
    else:
        row_cells = table.add_row().cells
        row_cells[0].text = "No action items identified."
        row_cells[1].text = "—"
        row_cells[2].text = "—"
        row_cells[3].text = "—"

    # --- Participants ---
    _add_heading(doc, "Participants")
    participants = analysis.get("participants") or []
    doc.add_paragraph(", ".join(participants) if participants else "Not identifiable from transcript.")

    # --- Sentiment ---
    _add_heading(doc, "Sentiment")
    sentiment = analysis.get("sentiment") or {}
    sentiment_line = sentiment.get("overall", "Unknown")
    if sentiment.get("explanation"):
        sentiment_line += f" — {sentiment['explanation']}"
    doc.add_paragraph(sentiment_line)

    # --- Topics ---
    _add_heading(doc, "Topics Discussed")
    topics = analysis.get("topics") or []
    if topics:
        for t in topics:
            name = t.get("topic", "Untitled") if isinstance(t, dict) else str(t)
            desc = t.get("description", "") if isinstance(t, dict) else ""
            line = name + (f" — {desc}" if desc else "")
            doc.add_paragraph(line, style="List Bullet")
    else:
        doc.add_paragraph("No topics identified.")

    # --- Meeting Insights ---
    _add_heading(doc, "Meeting Insights")
    insight_table = doc.add_table(rows=0, cols=2)
    insight_table.style = "Light List Accent 1"
    insight_rows = [
        ("Duration", str(meeting_insights.get("duration_seconds", "Unknown"))),
        ("Word Count", str(meeting_insights.get("word_count", 0))),
        ("Speaking Rate (WPM)", str(meeting_insights.get("speaking_rate_wpm", 0))),
        ("Estimated Speakers", f"~{meeting_insights.get('estimated_speakers', 1)} (approximate)"),
    ]
    for label, value in insight_rows:
        row_cells = insight_table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = value

    doc.save(output_path)
    log.info("DOCX report written to %s", output_path)
    return output_path
